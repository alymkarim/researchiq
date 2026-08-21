import io
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors


def export_analysis_pdf(
    title: str,
    authors: str | None,
    analysis: dict,
) -> bytes:
    """Export analysis as PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=8,
        textColor=colors.HexColor("#6c5ce7"),
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
    )

    story = []

    story.append(Paragraph(f"Analysis: {title}", title_style))
    if authors:
        story.append(Paragraph(f"Authors: {authors}", body_style))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        body_style,
    ))
    story.append(Spacer(1, 0.3 * inch))

    mode = analysis.get("analysis_mode", "unknown")
    story.append(Paragraph(f"Analysis Mode: {mode.upper()}", body_style))
    story.append(Spacer(1, 0.2 * inch))

    sections = [
        ("Executive Summary", analysis.get("summary", "")),
        ("Research Objective", analysis.get("objective", "")),
        ("Methodology", analysis.get("methodology", "")),
        ("Dataset / Sample", analysis.get("dataset", "")),
        ("Main Findings", analysis.get("findings", "")),
        ("Strengths", analysis.get("strengths", "")),
        ("Limitations", analysis.get("limitations", "")),
    ]

    for heading, content in sections:
        if content and content != "Not clearly stated":
            story.append(Paragraph(heading, heading_style))
            story.append(Paragraph(str(content), body_style))
            story.append(Spacer(1, 0.15 * inch))

    keywords = analysis.get("keywords", "")
    if keywords and keywords != "Not clearly stated":
        story.append(Paragraph("Keywords", heading_style))
        story.append(Paragraph(str(keywords), body_style))

    doc.build(story)
    return buffer.getvalue()


def export_analysis_docx(
    title: str,
    authors: str | None,
    analysis: dict,
) -> bytes:
    """Export analysis as DOCX."""
    doc = DocxDocument()

    doc.add_heading(f"Analysis: {title}", level=0)

    if authors:
        p = doc.add_paragraph()
        p.add_run("Authors: ").bold = True
        p.add_run(authors)

    p = doc.add_paragraph()
    p.add_run("Generated: ").bold = True
    p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))

    mode = analysis.get("analysis_mode", "unknown")
    p = doc.add_paragraph()
    p.add_run("Analysis Mode: ").bold = True
    p.add_run(mode.upper())

    doc.add_paragraph()

    sections = [
        ("Executive Summary", analysis.get("summary", "")),
        ("Research Objective", analysis.get("objective", "")),
        ("Methodology", analysis.get("methodology", "")),
        ("Dataset / Sample", analysis.get("dataset", "")),
        ("Main Findings", analysis.get("findings", "")),
        ("Strengths", analysis.get("strengths", "")),
        ("Limitations", analysis.get("limitations", "")),
    ]

    for heading, content in sections:
        if content and content != "Not clearly stated":
            doc.add_heading(heading, level=1)
            doc.add_paragraph(str(content))

    keywords = analysis.get("keywords", "")
    if keywords and keywords != "Not clearly stated":
        doc.add_heading("Keywords", level=1)
        doc.add_paragraph(str(keywords))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_comparison_pdf(
    documents: list[dict],
    comparison: dict,
) -> bytes:
    """Export comparison as PDF."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=18,
        spaceAfter=12,
    )
    heading_style = ParagraphStyle(
        "CustomHeading",
        parent=styles["Heading2"],
        fontSize=14,
        spaceAfter=8,
        textColor=colors.HexColor("#6c5ce7"),
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=10,
        spaceAfter=6,
    )

    story = []

    story.append(Paragraph("Paper Comparison", title_style))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        body_style,
    ))
    story.append(Spacer(1, 0.3 * inch))

    story.append(Paragraph("Compared Papers", heading_style))
    for doc_info in documents:
        story.append(Paragraph(
            f"• {doc_info.get('title', 'Untitled')}",
            body_style,
        ))
    story.append(Spacer(1, 0.2 * inch))

    for key, value in comparison.items():
        if key in ("keywords", "documents", "papers", "shared_keywords", "similarities"):
            continue
        if value:
            story.append(Paragraph(key.replace("_", " ").title(), heading_style))
            story.append(Paragraph(str(value), body_style))
            story.append(Spacer(1, 0.15 * inch))

    doc.build(story)
    return buffer.getvalue()
